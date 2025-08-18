#!/usr/bin/env python3

import sys
import json
import os
from pathlib import Path

def parse_pdf(file_path):
    """Parse PDF file and extract text"""
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                "text": text.strip(),
                "metadata": {
                    "pageCount": len(pdf_reader.pages),
                    "wordCount": len(text.split())
                }
            }
    except ImportError:
        raise Exception("PyPDF2 not installed. Run: pip install PyPDF2")
    except Exception as e:
        raise Exception(f"Failed to parse PDF: {str(e)}")

def parse_docx(file_path):
    """Parse DOCX file and extract text"""
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return {
            "text": text.strip(),
            "metadata": {
                "paragraphs": len(doc.paragraphs),
                "wordCount": len(text.split())
            }
        }
    except ImportError:
        raise Exception("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise Exception(f"Failed to parse DOCX: {str(e)}")

def main():
    if len(sys.argv) != 3:
        print(json.dumps({"error": "Usage: python document_parser.py <file_path> <file_type>"}))
        sys.exit(1)
    
    file_path = sys.argv[1]
    file_type = sys.argv[2].lower()
    
    if not os.path.exists(file_path):
        print(json.dumps({"error": f"File not found: {file_path}"}))
        sys.exit(1)
    
    try:
        if file_type == 'pdf':
            result = parse_pdf(file_path)
        elif file_type == 'docx':
            result = parse_docx(file_path)
        else:
            raise Exception(f"Unsupported file type: {file_type}")
        
        print(json.dumps(result))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)

if __name__ == "__main__":
    main()
